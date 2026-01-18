# services/export_service.py
"""
Сервис для экспорта документов в различные форматы.
Поддерживает экспорт в PDF, TXT, DOCX.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
from typing import Optional
from datetime import datetime

# Настраиваем логирование
logger = logging.getLogger(__name__)


class ExportService:
    """
    Сервис для экспорта документов в различные форматы.
    """

    def __init__(self, temp_folder: str):
        """
        Инициализация сервиса экспорта.

        Args:
            temp_folder: временная папка для создания файлов экспорта
        """
        self.temp_folder = temp_folder

        # Создаем временную папку, если не существует
        os.makedirs(temp_folder, exist_ok=True)

        logger.info(f"ExportService инициализирован: temp_folder={temp_folder}")

    def export_to_txt(
        self, content: str, filename: str, title: Optional[str] = None
    ) -> Optional[str]:
        """
        Экспортирует текст в TXT файл.

        Args:
            content: содержимое документа
            filename: имя файла (без расширения)
            title: заголовок документа (опционально)

        Returns:
            Путь к созданному файлу или None при ошибке
        """
        try:
            logger.info(f"Экспорт в TXT: {filename}")

            # Формируем путь к выходному файлу
            output_path = os.path.join(self.temp_folder, f"{filename}.txt")

            # Открываем файл для записи с кодировкой UTF-8 (для кириллицы)
            with open(output_path, "w", encoding="utf-8") as f:
                # Если есть заголовок, добавляем его
                if title:
                    f.write(f"{title}\n")
                    f.write("=" * len(title) + "\n\n")

                # Добавляем дату экспорта
                export_date = datetime.now().strftime("%d.%m.%Y %H:%M")
                f.write(f"Экспортировано: {export_date}\n\n")
                f.write("-" * 50 + "\n\n")

                # Записываем содержимое
                f.write(content)

            logger.info(f"TXT файл успешно создан: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Ошибка при экспорте в TXT: {str(e)}", exc_info=True)
            return None

    def export_to_docx(
        self, content: str, filename: str, title: Optional[str] = None
    ) -> Optional[str]:
        """
        Экспортирует текст в DOCX файл с форматированием.

        Args:
            content: содержимое документа
            filename: имя файла (без расширения)
            title: заголовок документа (опционально)

        Returns:
            Путь к созданному файлу или None при ошибке
        """
        try:
            logger.info(f"Экспорт в DOCX: {filename}")

            # Создаем новый документ Word
            doc = Document()

            # Настраиваем поля документа (1 дюйм со всех сторон)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Если есть заголовок, добавляем его
            if title:
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавляем информацию об экспорте
            export_info = doc.add_paragraph()
            export_date = datetime.now().strftime("%d.%m.%Y %H:%M")
            export_info.add_run(f"Экспортировано: {export_date}").italic = True
            export_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Добавляем разделитель
            doc.add_paragraph("_" * 50)

            # Добавляем пустую строку
            doc.add_paragraph()

            # Разбиваем содержимое на абзацы и добавляем в документ
            paragraphs = content.split("\n\n")

            for para_text in paragraphs:
                if para_text.strip():
                    # Создаем новый абзац
                    paragraph = doc.add_paragraph()

                    # Настраиваем стиль
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5
                    paragraph_format.space_after = Pt(6)

                    # Добавляем текст
                    run = paragraph.add_run(para_text.strip())
                    run.font.size = Pt(12)
                    run.font.name = "Times New Roman"

            # Формируем путь к выходному файлу
            output_path = os.path.join(self.temp_folder, f"{filename}.docx")

            # Сохраняем документ
            doc.save(output_path)

            logger.info(f"DOCX файл успешно создан: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Ошибка при экспорте в DOCX: {str(e)}", exc_info=True)
            return None

    def export_to_pdf(
        self, content: str, filename: str, title: Optional[str] = None, pdf_service=None
    ) -> Optional[str]:
        """
        Экспортирует текст в PDF файл.

        Args:
            content: содержимое документа
            filename: имя файла (без расширения)
            title: заголовок документа (опционально)
            pdf_service: экземпляр PDFService для создания PDF

        Returns:
            Путь к созданному файлу или None при ошибке
        """
        try:
            logger.info(f"Экспорт в PDF: {filename}")

            if not pdf_service:
                logger.error("PDFService не предоставлен")
                return None

            # Формируем путь к выходному файлу
            output_path = os.path.join(self.temp_folder, f"{filename}.pdf")

            # Используем PDFService для создания PDF
            success = pdf_service.create_pdf_from_text(
                text=content, output_pdf_path=output_path, title=title or filename
            )

            if success:
                logger.info(f"PDF файл успешно создан: {output_path}")
                return output_path
            else:
                logger.error("Не удалось создать PDF файл")
                return None

        except Exception as e:
            logger.error(f"Ошибка при экспорте в PDF: {str(e)}", exc_info=True)
            return None

    def cleanup_old_exports(self, max_age_seconds: int = 3600):
        """
        Удаляет старые файлы экспорта из временной папки.

        Args:
            max_age_seconds: максимальный возраст файла в секундах (по умолчанию 1 час)
        """
        try:
            logger.info(
                f"Очистка старых файлов экспорта (старше {max_age_seconds} секунд)"
            )

            current_time = datetime.now().timestamp()
            deleted_count = 0

            # Проходим по всем файлам во временной папке
            for filename in os.listdir(self.temp_folder):
                file_path = os.path.join(self.temp_folder, filename)

                # Проверяем, что это файл (не директория)
                if os.path.isfile(file_path):
                    # Получаем время последней модификации
                    file_age = current_time - os.path.getmtime(file_path)

                    # Если файл старше максимального возраста, удаляем
                    if file_age > max_age_seconds:
                        os.remove(file_path)
                        deleted_count += 1
                        logger.debug(f"Удален старый файл: {filename}")

            logger.info(f"Очистка завершена. Удалено файлов: {deleted_count}")

        except Exception as e:
            logger.error(f"Ошибка при очистке старых файлов: {str(e)}", exc_info=True)

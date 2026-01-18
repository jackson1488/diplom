# routes/documents.py
"""
Маршруты для управления документами.
Включает просмотр библиотеки, создание/удаление папок, управление документами.
"""
from flask import send_file  # Если еще нет

from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    jsonify,
    send_file,
    current_app,
)
from flask_login import current_user
import os
import logging
import shutil

from models import db
from models.document import Document
from models.folder import Folder
from utils.decorators import login_required
from utils.validators import validate_folder_name, validate_document_title
from utils.helpers import format_file_size, format_date
from services.document_service import DocumentService
from services.export_service import ExportService
from services.pdf_service import PDFService

# Настраиваем логирование
logger = logging.getLogger(__name__)

# Создаем blueprint для маршрутов документов
documents_bp = Blueprint("documents", __name__)


@documents_bp.route("/library")
@login_required
def library():
    """
    Главная страница библиотеки документов.
    Отображает все документы пользователя с возможностью фильтрации и сортировки.
    """
    # Получаем параметры из query string
    folder_id = request.args.get("folder_id", type=int)
    sort_by = request.args.get("sort_by", "created_at")
    order = request.args.get("order", "desc")
    search_query = request.args.get("q", "").strip()

    logger.info(
        f"Библиотека: user_id={current_user.id}, folder_id={folder_id}, sort={sort_by}"
    )

    # Базовый запрос документов пользователя
    query = Document.query.filter_by(user_id=current_user.id, is_archived=False)

    # Фильтр по папке
    if folder_id:
        query = query.filter_by(folder_id=folder_id)
        current_folder = Folder.query.filter_by(
            id=folder_id, user_id=current_user.id
        ).first()
    else:
        current_folder = None

    # Поиск по названию и содержимому
    if search_query:
        query = query.filter(
            db.or_(
                Document.title.contains(search_query),
                Document.content.contains(search_query),
                Document.description.contains(search_query),
            )
        )

    # Сортировка
    valid_sort_fields = ["created_at", "updated_at", "title", "file_size"]
    if sort_by in valid_sort_fields:
        order_column = getattr(Document, sort_by)
        if order == "desc":
            query = query.order_by(order_column.desc())
        else:
            query = query.order_by(order_column.asc())

    # Выполняем запрос
    documents = query.all()

    # Получаем все папки пользователя
    folders = (
        Folder.query.filter_by(user_id=current_user.id).order_by(Folder.name).all()
    )

    # Статистика
    stats = {
        "total_documents": Document.query.filter_by(
            user_id=current_user.id, is_archived=False
        ).count(),
        "total_folders": len(folders),
        "total_size": sum([doc.file_size or 0 for doc in documents]),
    }

    return render_template(
        "documents/library.html",
        documents=documents,
        folders=folders,
        current_folder=current_folder,
        stats=stats,
        sort_by=sort_by,
        order=order,
        search_query=search_query,
    )


@documents_bp.route("/view/<int:document_id>")
@login_required
def view_document(document_id):
    """
    Просмотр отдельного документа.
    Отображает метаданные, содержимое и миниатюру документа.
    """
    # Получаем документ
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()

    if not document:
        flash("Документ не найден", "danger")
        return redirect(url_for("documents.library"))

    # Обновляем время последнего просмотра
    document.update_last_viewed()

    logger.info(f"Просмотр документа: doc_id={document_id}, user_id={current_user.id}")

    return render_template("documents/view.html", document=document)


@documents_bp.route("/delete/<int:document_id>", methods=["POST"])
@login_required
def delete_document(document_id):
    """
    Удаление документа.
    Удаляет документ и все связанные файлы.
    """
    from flask import current_app

    # Инициализируем сервис документов
    doc_service = DocumentService(
        upload_folder=current_app.config["UPLOAD_FOLDER"],
        allowed_extensions=current_app.config["ALLOWED_EXTENSIONS"],
    )

    # Удаляем документ
    success = doc_service.delete_document(document_id, current_user.id)

    if success:
        flash("Документ успешно удален", "success")
    else:
        flash("Ошибка при удалении документа", "danger")

    return redirect(url_for("documents.library"))


@documents_bp.route("/update/<int:document_id>", methods=["POST"])
@login_required
def update_document(document_id):
    """
    Обновление метаданных документа.
    Изменяет название, описание, теги и другие поля.
    """
    # Получаем документ
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()

    if not document:
        return jsonify({"success": False, "error": "Документ не найден"}), 404

    # Получаем данные из формы
    title = request.form.get("title", "").strip()
    description = request.form.get("description", "").strip()
    tags = request.form.get("tags", "").strip()

    # Валидация
    if title:
        is_valid, error_msg = validate_document_title(title)
        if not is_valid:
            return jsonify({"success": False, "error": error_msg}), 400
        document.title = title

    if description:
        document.description = description

    if tags:
        document.tags = tags

    try:
        db.session.commit()
        logger.info(f"Документ обновлен: doc_id={document_id}")

        if request.is_json:
            return jsonify({"success": True, "message": "Документ обновлен"})
        else:
            flash("Документ успешно обновлен", "success")
            return redirect(url_for("documents.view_document", document_id=document_id))

    except Exception as e:
        db.session.rollback()
        logger.error(f"Ошибка при обновлении документа: {str(e)}", exc_info=True)

        if request.is_json:
            return jsonify({"success": False, "error": "Ошибка при обновлении"}), 500
        else:
            flash("Ошибка при обновлении документа", "danger")
            return redirect(url_for("documents.view_document", document_id=document_id))


@documents_bp.route("/move/<int:document_id>", methods=["POST"])
@login_required
def move_document(document_id):
    """
    Перемещение документа в другую папку.
    """
    from flask import current_app

    # Получаем целевую папку
    folder_id = request.form.get("folder_id", type=int)

    # Инициализируем сервис документов
    doc_service = DocumentService(
        upload_folder=current_app.config["UPLOAD_FOLDER"],
        allowed_extensions=current_app.config["ALLOWED_EXTENSIONS"],
    )

    # Перемещаем документ
    success = doc_service.move_document_to_folder(
        document_id, folder_id, current_user.id
    )

    if success:
        flash("Документ успешно перемещен", "success")
    else:
        flash("Ошибка при перемещении документа", "danger")

    return redirect(url_for("documents.library"))


@documents_bp.route("/toggle_favorite/<int:document_id>", methods=["POST"])
@login_required
def toggle_favorite(document_id):
    """
    Переключение статуса избранного документа.
    """
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()

    if not document:
        return jsonify({"success": False, "error": "Документ не найден"}), 404

    # Переключаем статус
    document.is_favorite = not document.is_favorite

    try:
        db.session.commit()
        logger.info(
            f"Статус избранного изменен: doc_id={document_id}, is_favorite={document.is_favorite}"
        )
        return jsonify({"success": True, "is_favorite": document.is_favorite})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Ошибка при изменении статуса: {str(e)}", exc_info=True)
        return jsonify({"success": False, "error": "Ошибка при обновлении"}), 500


@documents_bp.route("/download/<int:document_id>")
@login_required
def download_document(document_id):
    """
    Скачивание документа в оригинальном формате.
    """
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()

    if not document:
        flash("Документ не найден", "danger")
        return redirect(url_for("documents.library"))

    # Проверяем существование файла
    if not os.path.exists(document.file_path):
        flash("Файл документа не найден на сервере", "danger")
        return redirect(url_for("documents.view_document", document_id=document_id))

    logger.info(
        f"Скачивание документа: doc_id={document_id}, user_id={current_user.id}"
    )

    # Отправляем файл
    return send_file(
        document.file_path, as_attachment=True, download_name=document.original_filename
    )


@documents_bp.route("/export/<int:document_id>", methods=["GET", "POST"])
@login_required
def export_document(document_id):
    """Экспорт документа в различных форматах"""
    try:
        document = Document.query.filter_by(
            id=document_id, user_id=current_user.id
        ).first_or_404()

        # Получаем формат из POST или GET параметров
        if request.method == "POST":
            export_format = request.form.get("format", "pdf")
        else:  # GET
            export_format = request.args.get("format", "pdf")

        # Создаем временную папку
        export_folder = current_app.config["EXPORT_TEMP_FOLDER"]
        os.makedirs(export_folder, exist_ok=True)

        # Генерируем имя файла
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{document.title}_{timestamp}"

        if export_format == "pdf":
            # Экспорт в PDF
            output_path = os.path.join(export_folder, f"{filename}.pdf")

            # Если документ уже PDF - копируем
            if document.file_extension == ".pdf":
                import shutil

                shutil.copy2(document.file_path, output_path)
            else:
                # Создаем PDF из текста
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                c = canvas.Canvas(output_path, pagesize=letter)

                # Заголовок
                c.setFont("Helvetica-Bold", 16)
                c.drawString(50, 750, document.title)

                # Текст
                c.setFont("Helvetica", 12)
                text = document.content or document.ocr_text or "Нет текста"
                y = 700

                for line in text.split("\n"):
                    if y < 50:
                        c.showPage()
                        c.setFont("Helvetica", 12)
                        y = 750

                    # Обрезаем длинные строки
                    if len(line) > 80:
                        line = line[:80]

                    try:
                        c.drawString(50, y, line)
                    except:
                        # Если символы не поддерживаются, заменяем
                        c.drawString(50, y, line.encode("ascii", "ignore").decode())

                    y -= 15

                c.save()

            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{filename}.pdf",
                mimetype="application/pdf",
            )

        elif export_format == "docx":
            # Экспорт в DOCX
            output_path = os.path.join(export_folder, f"{filename}.docx")

            from docx import Document as DocxDocument

            doc = DocxDocument()
            doc.add_heading(document.title, 0)

            text = document.content or document.ocr_text or "Нет текста"
            doc.add_paragraph(text)

            doc.save(output_path)

            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{filename}.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        elif export_format == "txt":
            # Экспорт в TXT
            output_path = os.path.join(export_folder, f"{filename}.txt")

            text = document.content or document.ocr_text or "Нет текста"

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"{document.title}\n\n")
                f.write(text)

            return send_file(
                output_path,
                as_attachment=True,
                download_name=f"{filename}.txt",
                mimetype="text/plain",
            )

        else:
            flash("Неподдерживаемый формат экспорта", "danger")
            return redirect(url_for("documents.view_document", document_id=document_id))

    except Exception as e:
        logger.error(f"Ошибка экспорта документа: {e}", exc_info=True)
        flash("Ошибка при экспорте документа", "danger")
        return redirect(url_for("documents.view_document", document_id=document_id))

    logger.info(f"Экспорт документа: doc_id={document_id}, format={format}")

    # Отправляем файл
    return send_file(
        export_path, as_attachment=True, download_name=f"{safe_filename}.{format}"
    )


# === УПРАВЛЕНИЕ ПАПКАМИ ===


@documents_bp.route("/folders")
@login_required
def folders():
    """
    Страница управления папками.
    Отображает все папки пользователя.
    """
    folders = (
        Folder.query.filter_by(user_id=current_user.id).order_by(Folder.name).all()
    )

    return render_template("documents/folders.html", folders=folders)


@documents_bp.route("/folder/create", methods=["POST"])
@login_required
def create_folder():
    """
    Создание новой папки.
    """
    # Получаем данные из формы
    name = request.form.get("name", "").strip()
    description = request.form.get("description", "").strip()
    color = request.form.get("color", "#3498db").strip()

    # Валидация названия
    is_valid, error_msg = validate_folder_name(name)
    if not is_valid:
        flash(error_msg, "danger")
        return redirect(url_for("documents.folders"))

    # Проверяем уникальность названия для пользователя
    existing_folder = Folder.query.filter_by(user_id=current_user.id, name=name).first()
    if existing_folder:
        flash("Папка с таким названием уже существует", "warning")
        return redirect(url_for("documents.folders"))

    try:
        # Создаем новую папку
        new_folder = Folder(
            name=name, description=description, color=color, user_id=current_user.id
        )

        db.session.add(new_folder)
        db.session.commit()

        logger.info(
            f"Создана папка: folder_id={new_folder.id}, user_id={current_user.id}"
        )
        flash("Папка успешно создана", "success")

    except Exception as e:
        db.session.rollback()
        logger.error(f"Ошибка при создании папки: {str(e)}", exc_info=True)
        flash("Ошибка при создании папки", "danger")

    return redirect(url_for("documents.folders"))


@documents_bp.route("/folder/update/<int:folder_id>", methods=["POST"])
@login_required
def update_folder(folder_id):
    """
    Обновление папки.
    """
    folder = Folder.query.filter_by(id=folder_id, user_id=current_user.id).first()

    if not folder:
        flash("Папка не найдена", "danger")
        return redirect(url_for("documents.folders"))

    # Получаем данные из формы
    name = request.form.get("name", "").strip()
    description = request.form.get("description", "").strip()
    color = request.form.get("color", "").strip()

    # Валидация
    if name:
        is_valid, error_msg = validate_folder_name(name)
        if not is_valid:
            flash(error_msg, "danger")
            return redirect(url_for("documents.folders"))
        folder.name = name

    if description:
        folder.description = description

    if color:
        folder.color = color

    try:
        db.session.commit()
        logger.info(f"Папка обновлена: folder_id={folder_id}")
        flash("Папка успешно обновлена", "success")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Ошибка при обновлении папки: {str(e)}", exc_info=True)
        flash("Ошибка при обновлении папки", "danger")

    return redirect(url_for("documents.folders"))


@documents_bp.route("/folder/delete/<int:folder_id>", methods=["POST"])
@login_required
def delete_folder(folder_id):
    """
    Удаление папки.
    Документы из папки не удаляются, а перемещаются в корень.
    """
    folder = Folder.query.filter_by(id=folder_id, user_id=current_user.id).first()

    if not folder:
        flash("Папка не найдена", "danger")
        return redirect(url_for("documents.folders"))

    try:
        # Перемещаем все документы из папки в корень
        documents = Document.query.filter_by(folder_id=folder_id).all()
        for doc in documents:
            doc.folder_id = None

        # Удаляем папку
        db.session.delete(folder)
        db.session.commit()

        logger.info(
            f"Папка удалена: folder_id={folder_id}, перемещено документов: {len(documents)}"
        )
        flash(
            f"Папка удалена. {len(documents)} документов перемещено в корень", "success"
        )

    except Exception as e:
        db.session.rollback()
        logger.error(f"Ошибка при удалении папки: {str(e)}", exc_info=True)
        flash("Ошибка при удалении папки", "danger")

    return redirect(url_for("documents.folders"))
